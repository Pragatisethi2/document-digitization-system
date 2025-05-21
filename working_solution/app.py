import os
from flask import Flask, render_template, request, jsonify, send_file
import google.generativeai as genai
# from dotenv import load_dotenv
import pandas as pd
from werkzeug.utils import secure_filename
import base64
from docx import Document
from docx.shared import Inches
from fpdf import FPDF
import json
from datetime import datetime
import traceback
import logging

# # Load environment variables
# load_dotenv()

# Configure logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

# Configure Gemini API
genai.configure(api_key='AIzaSyCxCJoOU1A5JPDAwtmpt5nr-Q97jTqLNzg')
model = genai.GenerativeModel('gemini-1.5-flash')

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['EXPORT_FOLDER'] = 'exports'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Ensure required folders exist
for folder in [app.config['UPLOAD_FOLDER'], app.config['EXPORT_FOLDER']]:
    os.makedirs(folder, exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in {'png', 'jpg', 'jpeg', 'pdf'}

def get_file_data(file_path):
    try:
        with open(file_path, 'rb') as file:
            data = file.read()
        return base64.b64encode(data).decode('utf-8')
    except Exception as e:
        logger.error(f"Error reading file {file_path}: {str(e)}")
        raise

def export_to_txt(data, filename):
    """Export data to a formatted text file"""
    try:
        filepath = os.path.join(app.config['EXPORT_FOLDER'], f"{filename}.txt")
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write("Digitized Document Data\n")
            f.write("=" * 50 + "\n\n")
            for key, value in data.items():
                f.write(f"{key}: {value}\n")
        return filepath
    except Exception as e:
        logger.error(f"Error creating TXT file: {str(e)}")
        raise

def export_to_docx(data, filename):
    """Export data to a Word document with table"""
    try:
        doc = Document()
        doc.add_heading('Digitized Document Data', 0)
        
        # Add timestamp
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Add table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Field'
        header_cells[1].text = 'Value'
        
        for key, value in data.items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(key)
            row_cells[1].text = str(value)

        filepath = os.path.join(app.config['EXPORT_FOLDER'], f"{filename}.docx")
        doc.save(filepath)
        return filepath
    except Exception as e:
        logger.error(f"Error creating DOCX file: {str(e)}")
        raise

def export_to_pdf(data, filename):
    """Export data to a PDF with table"""
    try:
        pdf = FPDF()
        pdf.add_page()
        
        # Set font
        pdf.set_font("Arial", size=12)
        
        # Add title
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(200, 10, txt="Digitized Document Data", ln=True, align='C')
        pdf.ln(10)
        
        # Add timestamp
        pdf.set_font("Arial", size=10)
        pdf.cell(200, 10, txt=f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ln=True, align='L')
        pdf.ln(10)
        
        # Set up table
        pdf.set_font("Arial", 'B', 12)
        col_width = pdf.get_string_width("Field Name    ")
        value_width = 190 - col_width
        
        # Add table headers
        pdf.cell(col_width, 10, "Field", 1, 0, 'L')
        pdf.cell(value_width, 10, "Value", 1, 1, 'L')
        
        # Add table data
        pdf.set_font("Arial", size=12)
        for key, value in data.items():
            # Handle multi-line content
            key_str = str(key)
            value_str = str(value)
            
            pdf.cell(col_width, 10, key_str, 1, 0, 'L')
            pdf.cell(value_width, 10, value_str, 1, 1, 'L')

        filepath = os.path.join(app.config['EXPORT_FOLDER'], f"{filename}.pdf")
        pdf.output(filepath)
        return filepath
    except Exception as e:
        logger.error(f"Error creating PDF file: {str(e)}")
        raise

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    try:
        if 'file' not in request.files:
            logger.error("No file part in request")
            return jsonify({'error': 'No file part'}), 400
        
        file = request.files['file']
        fields = request.form.get('fields', '').split(',')
        
        if file.filename == '':
            logger.error("No selected file")
            return jsonify({'error': 'No selected file'}), 400
        
        if not fields:
            logger.error("No fields specified")
            return jsonify({'error': 'No fields specified'}), 400
        
        if file and allowed_file(file.filename):
            # Save uploaded file
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            logger.info(f"File saved: {filepath}")
            
            try:
                # Process with Gemini
                file_data = get_file_data(filepath)
                logger.info("File data encoded successfully")
                
                prompt = f"""Please extract the following information from this document: {', '.join(fields)}. 
                Format the response as a JSON object with these fields as keys."""
                
                logger.info("Sending request to Gemini API")
                response = model.generate_content([
                    prompt,
                    {'mime_type': 'image/jpeg', 'data': file_data}
                ])
                logger.info("Received response from Gemini API")
                
                try:
                    # Parse the response
                    response_text = response.text.strip()
                    logger.debug(f"Raw API response: {response_text}")
                    
                    # Try to find JSON-like content in the response
                    start_idx = response_text.find('{')
                    end_idx = response_text.rfind('}')
                    
                    if start_idx != -1 and end_idx != -1:
                        json_str = response_text[start_idx:end_idx + 1]
                        data = json.loads(json_str)
                    else:
                        raise ValueError("No valid JSON found in response")
                    
                    logger.info("Response parsed successfully")
                    
                    # Generate base filename without extension
                    base_filename = f"digitized_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                    
                    # Export in different formats
                    txt_path = export_to_txt(data, base_filename)
                    docx_path = export_to_docx(data, base_filename)
                    pdf_path = export_to_pdf(data, base_filename)
                    
                    # Create Excel file
                    df = pd.DataFrame([data])
                    excel_path = os.path.join(app.config['EXPORT_FOLDER'], f"{base_filename}.xlsx")
                    df.to_excel(excel_path, index=False)
                    
                    logger.info("All files exported successfully")
                    
                    return jsonify({
                        'message': 'Data extracted and saved successfully',
                        'data': data,
                        'exports': {
                            'txt': os.path.basename(txt_path),
                            'docx': os.path.basename(docx_path),
                            'pdf': os.path.basename(pdf_path),
                            'excel': os.path.basename(excel_path)
                        }
                    })
                except Exception as e:
                    logger.error(f"Error processing data: {str(e)}")
                    logger.error(traceback.format_exc())
                    return jsonify({'error': f'Error processing data: {str(e)}'}), 500
                    
            except Exception as e:
                logger.error(f"Error with Gemini API: {str(e)}")
                logger.error(traceback.format_exc())
                return jsonify({'error': f'Error with Gemini API: {str(e)}'}), 500
            
        logger.error("Invalid file type")
        return jsonify({'error': 'Invalid file type'}), 400
        
    except Exception as e:
        logger.error(f"Unexpected error: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': f'An unexpected error occurred: {str(e)}'}), 500

@app.route('/download/<filename>')
def download_file(filename):
    """Route to download exported files"""
    try:
        return send_file(
            os.path.join(app.config['EXPORT_FOLDER'], filename),
            as_attachment=True
        )
    except Exception as e:
        logger.error(f"Error downloading file: {str(e)}")
        return jsonify({'error': f'Error downloading file: {str(e)}'}), 404

if __name__ == '__main__':
    app.run(debug=True) 